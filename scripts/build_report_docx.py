from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = PROJECT_ROOT / "report" / "Report.md"
DEFAULT_OUTPUT = PROJECT_ROOT / "report" / "Report.docx"
FIGURES_PATH = PROJECT_ROOT / "outputs" / "figures"

PRIMARY = "123C43"
SECONDARY = "D8A31A"
LIGHT = "F5F2E8"
PALE_BLUE = "EAF3F4"
ROW_ALT = "F7F7F7"

AUTHORS = [
    ("Cavaignac Romain", "TP1458"),
    ("Dubernet Mathieu", "TP145868"),
    ("Haegeman Victor", "TP145873"),
]

FIGURES = [
    ("assignment_pipeline_overview.png", "Figure 1. End-to-end project pipeline from raw data sources to validated assignment outputs."),
    ("target_distribution.png", "Figure 2. Target distribution for the binary top-10 classification task."),
    ("rows_by_season.png", "Figure 3. Number of driver-race rows available per season."),
    ("top10_rate_by_season.png", "Figure 4. Top-10 target rate by season, used to check class stability over time."),
    ("grid_vs_finish.png", "Figure 5. Relationship between starting grid and final position."),
    ("algorithm_holdout_summary.png", "Figure 6. Holdout comparison of the main classification algorithms on the 2025 season."),
    ("model_metrics_table.png", "Figure 7. Validation metric table for the compared classification algorithms."),
    ("model_comparison.png", "Figure 8. Compact model comparison across F1, ROC-AUC and race precision@10."),
    ("algorithm_validation_summary.png", "Figure 9. Expanding-window validation showing model stability across seasons."),
    ("rolling_backtest.png", "Figure 10. Rolling backtest race precision@10 by model and test season."),
    ("feature_importance.png", "Figure 11. Most influential features for the selected top-10 classifier."),
    ("confusion_matrix.png", "Figure 12. Confusion matrix for the selected holdout classifier."),
    ("neural_network_tuning.png", "Figure 13. Neural-network classifier tuning results, included as an extension."),
    ("neural_network_embedding_3d.png", "Figure 14. Static view of the interactive 3D neural-network hidden-space embedding."),
    ("position_model_comparison.png", "Figure 15. Finish-position ranking model comparison, included as an extension."),
    ("predictions/model_precision_by_race.png", "Figure 16. Correct predicted top-10 drivers per race and model on the 2025 holdout season."),
    ("predictions/model_hit_heatmap.png", "Figure 17. Race-level heatmap of top-10 hits by model."),
    ("predictions/model_points_captured.png", "Figure 18. Average actual points captured by each model's predicted top 10."),
    ("predictions/race_overviews/2025_08_monaco_grand_prix.png", "Figure 19. Example race overview with real result, model scoreboard, consensus picks and predicted top 10 chips."),
    ("predictions/race_cards/2025_08_monaco_grand_prix.png", "Figure 20. Detailed race card showing actual podium, virtual podiums and predicted top 10 lists."),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build report/Report.docx from report/Report.md.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str | None = None,
    size: float = 9,
    alignment: int | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 120, start: int = 120, bottom: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        element = margins.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D7D7D7", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_cell_run(paragraph, text: str, *, size: float, color: str, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.different_first_page_header_footer = True

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(PRIMARY)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(PRIMARY)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor.from_string(PRIMARY)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(12)
    styles["List Number"].font.name = "Arial"
    styles["List Number"].font.size = Pt(12)

    footer = section.footer.paragraphs[0]
    footer.text = "F1 Top-10 Finish Prediction | CX016-2.5-3-IML"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(12)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_cover_page(document: Document) -> None:
    document.add_paragraph()

    header = document.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(header, PRIMARY, "0")
    header_cell = header.cell(0, 0)
    set_cell_shading(header_cell, PRIMARY)
    set_cell_margins(header_cell, top=360, start=260, bottom=340, end=260)
    header_paragraph = header_cell.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_cell_run(header_paragraph, "F1 Top-10 Finish Prediction", size=30, color="FFFFFF", bold=True)
    add_cell_run(header_paragraph, "\nMachine Learning Classification Report", size=17, color="FFFFFF")
    add_cell_run(header_paragraph, "\nCX016-2.5-3-IML - Introduction to Machine Learning", size=12, color="EAF3F4")

    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(accent, SECONDARY, "0")
    accent_cell = accent.cell(0, 0)
    set_cell_shading(accent_cell, SECONDARY)
    set_cell_margins(accent_cell, top=28, bottom=28)
    accent_cell.paragraphs[0].add_run("")

    document.add_paragraph()

    authors_table = document.add_table(rows=1 + len(AUTHORS), cols=2)
    authors_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    authors_table.style = "Table Grid"
    set_table_borders(authors_table, "D0D7D8", "6")
    set_cell_shading(authors_table.cell(0, 0), PRIMARY)
    set_cell_text(authors_table.cell(0, 0), "Group members", bold=True, color="FFFFFF", size=12)
    set_cell_shading(authors_table.cell(0, 1), PRIMARY)
    set_cell_text(authors_table.cell(0, 1), "TP number", bold=True, color="FFFFFF", size=12)
    for row_index, (name, tp_number) in enumerate(AUTHORS, start=1):
        fill = LIGHT if row_index % 2 else PALE_BLUE
        set_cell_shading(authors_table.cell(row_index, 0), fill)
        set_cell_text(authors_table.cell(row_index, 0), name, size=12)
        set_cell_shading(authors_table.cell(row_index, 1), fill)
        set_cell_text(authors_table.cell(row_index, 1), tp_number, size=12)
    for row in authors_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=110, start=140, bottom=110, end=140)

    document.add_paragraph()

    metadata = [
        ("Module code", "CX016-2.5-3-IML"),
        ("Module title", "Introduction to Machine Learning"),
        ("Class / intake code", "CSSE___CX016-2.5-3-IML-L-1___2026-01-30"),
        ("Hand out date", "13 February 2026"),
        ("Hand in date", "08 May 2026"),
        ("Weightage", "30%"),
        ("Deliverable", "Report, notebook, dataset, scripts and submission ZIP"),
        ("Main task", "Predict whether each Formula 1 driver finishes in the top 10"),
        ("Dataset", "6,999 driver-race rows, 201 variables, seasons 2010-2026"),
        ("Primary model", "Random Forest classifier"),
    ]
    table = document.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table, "D0D7D8", "6")
    for row_index, (label, value) in enumerate(metadata):
        set_cell_shading(table.cell(row_index, 0), PRIMARY)
        set_cell_text(table.cell(row_index, 0), label, bold=True, color="FFFFFF", size=12)
        set_cell_shading(table.cell(row_index, 1), LIGHT if row_index % 2 == 0 else PALE_BLUE)
        set_cell_text(table.cell(row_index, 1), value, size=12)
        set_cell_margins(table.cell(row_index, 0), top=95, start=130, bottom=95, end=130)
        set_cell_margins(table.cell(row_index, 1), top=95, start=130, bottom=95, end=130)

    document.add_paragraph()
    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_run = badge.add_run("Temporal validation: 2025 holdout | Dataset: 2010-2026 | Target: top10_finish")
    badge_run.font.name = "Arial"
    badge_run.font.size = Pt(12)
    badge_run.font.color.rgb = RGBColor.from_string("555555")

    document.add_page_break()


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(set(cell) <= {"-", ":"} and "-" in cell for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        rows.append([cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")])

    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, PRIMARY)
                set_cell_text(cell, value, bold=True, color="FFFFFF")
            else:
                set_cell_shading(cell, PALE_BLUE if row_index % 2 else ROW_ALT)
                set_cell_text(cell, value)
    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("333333")


def add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("# "):
        document.add_heading(stripped[2:], level=1)
    elif stripped.startswith("## "):
        document.add_heading(stripped[3:], level=2)
    elif stripped.startswith("### "):
        document.add_heading(stripped[4:], level=3)
    elif stripped.startswith("- "):
        document.add_paragraph(stripped[2:].replace("`", ""), style="List Bullet")
    elif len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:4]:
        document.add_paragraph(stripped.split(". ", 1)[1].replace("`", ""), style="List Number")
    else:
        paragraph = document.add_paragraph(stripped.replace("`", ""))
        paragraph.paragraph_format.line_spacing = 1.15


def add_figures(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Figures", level=2)
    for figure_name, caption in FIGURES:
        figure_path = FIGURES_PATH / figure_name
        if not figure_path.exists() or figure_path.stat().st_size == 0:
            continue

        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(9)
        caption_run.font.color.rgb = RGBColor.from_string("333333")
        document.add_picture(str(figure_path), width=Inches(6.2))


def add_markdown_body(document: Document, lines: list[str]) -> None:
    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_title = False

    while index < len(lines):
        line = lines[index]

        if not skipped_title and line.startswith("# "):
            skipped_title = True
            index += 1
            continue

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if is_table_line(line):
            table_lines = []
            while index < len(lines) and is_table_line(lines[index]):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        add_markdown_line(document, line)
        index += 1


def main() -> None:
    args = parse_args()
    input_path = args.input if args.input.is_absolute() else PROJECT_ROOT / args.input
    output_path = args.output if args.output.is_absolute() else PROJECT_ROOT / args.output

    if not input_path.exists():
        raise FileNotFoundError(f"Missing report markdown: {input_path}")

    document = Document()
    configure_document(document)
    add_cover_page(document)
    add_markdown_body(document, input_path.read_text(encoding="utf-8").splitlines())
    add_figures(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
